from __future__ import annotations

from collections import defaultdict
from io import BytesIO
from pathlib import Path
import os
import re
from typing import Iterable
from urllib.parse import quote

from django.db.models import Count, Q, QuerySet
from django.http import HttpResponse
from django.utils import timezone
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from drf_spectacular.types import OpenApiTypes
from drf_spectacular.utils import extend_schema
from rest_framework import permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView

from alumni.models import AcademicGroup, AlumniProfile
from employers.models import Employer


DEFAULT_REPORT_LANGUAGE = "ru"
REPORT_LANGUAGES = {"ru", "en", "kg"}
REPORT_LANGUAGE_ALIASES = {
    "ru": "ru",
    "ru-ru": "ru",
    "en": "en",
    "en-us": "en",
    "en-gb": "en",
    "kg": "kg",
    "ky": "kg",
    "ky-kg": "kg",
}

REPORT_TEXTS: dict[str, dict[str, str]] = {
    "ru": {
        "title": "Отчет трудоустройства кафедры «ИСТ им. акад. А. Жайнакова»",
        "subtitle": "Кафедра «Информационные системы и технологии имени академика А. Жайнакова», ИЭТ КГТУ им. И. Раззакова",
        "all_selection": "Все группы и годы выпуска",
        "section_1": "Раздел 1. Реестр выпускников и сведения о занятости",
        "section_2": "Раздел 2. Сводные показатели трудоустройства",
        "chair_signature": "Зав. кафедрой «ИСТ им. акад. А. Жайнакова» ____________________",
        "unsupported_export": "Неподдерживаемый формат экспорта. Используйте pdf, docx или xlsx.",
        "file_stem": "отчет_трудоустройство_выпускников",
        "sheet_graduates": "Выпускники",
        "sheet_summary": "Сводка",
        "number": "№",
        "full_name": "ФИО",
        "group": "Группа",
        "graduation_year": "Год выпуска",
        "year_graduation": "Год окончания",
        "specialty": "Специальность",
        "specialty_diploma": "Специальность по диплому",
        "employment": "Трудоустройство",
        "employed_specialty_header": "по спец.\n(указать должность, место работы)",
        "not_specialty_self_employed_header": "не по спец.\n/ самозанятость",
        "not_working_unknown_header": "не работает\n/ неизвестно",
        "continued_education_where": "Продолжил обучение\nУказать где именно",
        "useful_question": "Что из преподававшегося на факультете наиболее полезно в вашей сегодняшней работе?",
        "self_study_question": "Что пришлось изучать самостоятельно?",
        "employed_specialty": "По спец. (должность, место работы)",
        "not_specialty_self_employed": "Не по спец. / самозанятость",
        "not_working_unknown": "Не работает / неизвестно",
        "continued_education": "Продолжил обучение",
        "useful_work": "Полезно в работе",
        "self_study": "Изучал самостоятельно",
        "direction_profile": "Направление / профиль",
        "form_level": "Форма / уровень",
        "graduates_count": "Кол-во выпускников",
        "surveyed_count": "Кол-во опрошенных",
        "works_specialty_percent": "Работают по спец. (%)",
        "works_not_specialty_percent": "Работают не по спец. (%)",
        "self_employed_percent": "Самозанятые (%)",
        "continued_education_percent": "Продолжили учебу (%)",
        "not_working_percent": "Не работают (%)",
        "unknown_percent": "Неизвестно (%)",
        "year_filter": "Год выпуска",
        "group_filter": "Группа",
        "direction_filter": "Направление",
        "study_form_filter": "Форма обучения",
        "degree_level_filter": "Уровень",
        "status_filter": "Статус",
        "search_filter": "Поиск",
        "not_working": "Не работает",
        "unknown": "Неизвестно",
        "no_year": "Без года",
    },
    "en": {
        "title": "Graduate employment report of the IST Department named after Acad. A. Zhainakov",
        "subtitle": "Department of Information Systems and Technologies named after Academician A. Zhainakov, IET KSTU named after I. Razzakov",
        "all_selection": "All groups and graduation years",
        "section_1": "Section 1. Graduate register and employment details",
        "section_2": "Section 2. Summary employment indicators",
        "chair_signature": "Head of the IST Department named after Acad. A. Zhainakov ____________________",
        "unsupported_export": "Unsupported export format. Use pdf, docx or xlsx.",
        "file_stem": "graduate_employment_report",
        "sheet_graduates": "Graduates",
        "sheet_summary": "Summary",
        "number": "No.",
        "full_name": "Full name",
        "group": "Group",
        "graduation_year": "Graduation year",
        "year_graduation": "Graduation year",
        "specialty": "Specialty",
        "specialty_diploma": "Diploma specialty",
        "employment": "Employment",
        "employed_specialty_header": "by specialty\n(position, workplace)",
        "not_specialty_self_employed_header": "not by specialty\n/ self-employed",
        "not_working_unknown_header": "not working\n/ unknown",
        "continued_education_where": "Continued education\nplace",
        "useful_question": "What from the faculty curriculum is most useful in your current work?",
        "self_study_question": "What did you have to study independently?",
        "employed_specialty": "By specialty (position, workplace)",
        "not_specialty_self_employed": "Not by specialty / self-employed",
        "not_working_unknown": "Not working / unknown",
        "continued_education": "Continued education",
        "useful_work": "Useful in work",
        "self_study": "Studied independently",
        "direction_profile": "Direction / profile",
        "form_level": "Study form / level",
        "graduates_count": "Graduates",
        "surveyed_count": "Surveyed",
        "works_specialty_percent": "Works by specialty (%)",
        "works_not_specialty_percent": "Works not by specialty (%)",
        "self_employed_percent": "Self-employed (%)",
        "continued_education_percent": "Continued education (%)",
        "not_working_percent": "Not working (%)",
        "unknown_percent": "Unknown (%)",
        "year_filter": "Graduation year",
        "group_filter": "Group",
        "direction_filter": "Direction",
        "study_form_filter": "Study form",
        "degree_level_filter": "Level",
        "status_filter": "Status",
        "search_filter": "Search",
        "not_working": "Not working",
        "unknown": "Unknown",
        "no_year": "No year",
    },
    "kg": {
        "title": "Акад. А. Жайнаков атындагы ИСТ кафедрасынын бүтүрүүчүлөрүнүн жумушка орношуусу боюнча отчет",
        "subtitle": "Академик А. Жайнаков атындагы маалыматтык системалар жана технологиялар кафедрасы, И. Раззаков атындагы КМТУнун ЭТИ",
        "all_selection": "Бардык тайпалар жана бүтүрүү жылдары",
        "section_1": "1-бөлүм. Бүтүрүүчүлөрдүн реестри жана жумушка орношуу маалыматтары",
        "section_2": "2-бөлүм. Жумушка орношуу боюнча жыйынтык көрсөткүчтөр",
        "chair_signature": "Акад. А. Жайнаков атындагы ИСТ кафедрасынын башчысы ____________________",
        "unsupported_export": "Экспорт форматы колдоого алынбайт. pdf, docx же xlsx колдонуңуз.",
        "file_stem": "бүтүрүүчүлөр_жумушка_орношуу_отчету",
        "sheet_graduates": "Бүтүрүүчүлөр",
        "sheet_summary": "Жыйынтык",
        "number": "№",
        "full_name": "Ф.А.А.",
        "group": "Тайпа",
        "graduation_year": "Бүтүрүү жылы",
        "year_graduation": "Бүтүрүү жылы",
        "specialty": "Адистик",
        "specialty_diploma": "Диплом боюнча адистик",
        "employment": "Жумушка орношуу",
        "employed_specialty_header": "адистиги боюнча\n(кызматы, иштеген жери)",
        "not_specialty_self_employed_header": "адистиги боюнча эмес\n/ өз алдынча иш",
        "not_working_unknown_header": "иштебейт\n/ белгисиз",
        "continued_education_where": "Окуусун улантты\nкайда",
        "useful_question": "Факультетте окутулган кайсы нерселер азыркы ишиңизде эң пайдалуу болду?",
        "self_study_question": "Эмнени өз алдынча үйрөнүүгө туура келди?",
        "employed_specialty": "Адистиги боюнча (кызматы, иштеген жери)",
        "not_specialty_self_employed": "Адистиги боюнча эмес / өз алдынча иш",
        "not_working_unknown": "Иштебейт / белгисиз",
        "continued_education": "Окуусун улантты",
        "useful_work": "Иште пайдалуу",
        "self_study": "Өз алдынча үйрөнгөн",
        "direction_profile": "Багыт / профиль",
        "form_level": "Окуу формасы / деңгээли",
        "graduates_count": "Бүтүрүүчүлөрдүн саны",
        "surveyed_count": "Сурамжылангандар",
        "works_specialty_percent": "Адистиги боюнча иштейт (%)",
        "works_not_specialty_percent": "Адистиги боюнча эмес иштейт (%)",
        "self_employed_percent": "Өз алдынча иштегендер (%)",
        "continued_education_percent": "Окуусун уланткандар (%)",
        "not_working_percent": "Иштебегендер (%)",
        "unknown_percent": "Белгисиз (%)",
        "year_filter": "Бүтүрүү жылы",
        "group_filter": "Тайпа",
        "direction_filter": "Багыт",
        "study_form_filter": "Окуу формасы",
        "degree_level_filter": "Деңгээли",
        "status_filter": "Статус",
        "search_filter": "Издөө",
        "not_working": "Иштебейт",
        "unknown": "Белгисиз",
        "no_year": "Жылы жок",
    },
}

CHOICE_LABELS: dict[str, dict[str, dict[str, str]]] = {
    "study_form": {
        "ru": {AcademicGroup.StudyForm.FULL_TIME: "Очное", AcademicGroup.StudyForm.PART_TIME: "Заочное (ДОТ)"},
        "en": {AcademicGroup.StudyForm.FULL_TIME: "Full-time", AcademicGroup.StudyForm.PART_TIME: "Part-time / distance"},
        "kg": {AcademicGroup.StudyForm.FULL_TIME: "Күндүзгү", AcademicGroup.StudyForm.PART_TIME: "Сырттан / ДОТ"},
    },
    "degree_level": {
        "ru": {AcademicGroup.DegreeLevel.BACHELOR: "Бакалавриат", AcademicGroup.DegreeLevel.MASTER: "Магистратура"},
        "en": {AcademicGroup.DegreeLevel.BACHELOR: "Bachelor", AcademicGroup.DegreeLevel.MASTER: "Master"},
        "kg": {AcademicGroup.DegreeLevel.BACHELOR: "Бакалавриат", AcademicGroup.DegreeLevel.MASTER: "Магистратура"},
    },
    "employment_status": {
        "ru": {
            AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY: "Работает по специальности",
            AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY: "Работает не по специальности",
            AlumniProfile.EmploymentStatus.SELF_EMPLOYED: "Самозанятый / предприниматель",
            AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION: "Продолжает обучение",
            AlumniProfile.EmploymentStatus.UNEMPLOYED: "Не работает",
            AlumniProfile.EmploymentStatus.LOST_CONTACT: "Неизвестно",
        },
        "en": {
            AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY: "Works by specialty",
            AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY: "Works not by specialty",
            AlumniProfile.EmploymentStatus.SELF_EMPLOYED: "Self-employed / entrepreneur",
            AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION: "Continues education",
            AlumniProfile.EmploymentStatus.UNEMPLOYED: "Not working",
            AlumniProfile.EmploymentStatus.LOST_CONTACT: "Unknown",
        },
        "kg": {
            AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY: "Адистиги боюнча иштейт",
            AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY: "Адистиги боюнча эмес иштейт",
            AlumniProfile.EmploymentStatus.SELF_EMPLOYED: "Өз алдынча иштеген / ишкер",
            AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION: "Окуусун улантууда",
            AlumniProfile.EmploymentStatus.UNEMPLOYED: "Иштебейт",
            AlumniProfile.EmploymentStatus.LOST_CONTACT: "Белгисиз",
        },
    },
}


def _normalize_report_language(value: str | None) -> str:
    if not value:
        return DEFAULT_REPORT_LANGUAGE
    cleaned = value.lower().strip().replace("_", "-")
    first = cleaned.split(",", 1)[0].split(";", 1)[0].strip()
    return REPORT_LANGUAGE_ALIASES.get(first, REPORT_LANGUAGE_ALIASES.get(first[:2], DEFAULT_REPORT_LANGUAGE))


def _texts(lang: str) -> dict[str, str]:
    return REPORT_TEXTS.get(lang, REPORT_TEXTS[DEFAULT_REPORT_LANGUAGE])


def _choice_label(category: str, value: str | None, lang: str) -> str:
    if not value:
        return "—"
    return CHOICE_LABELS.get(category, {}).get(lang, {}).get(value, str(value))


class RotatedParagraph(Paragraph):
    """A paragraph rotated 90 degrees for narrow PDF table headers."""

    def wrap(self, avail_width, avail_height):
        width, height = Paragraph.wrap(self, avail_height, avail_width)
        self._rotated_width = height
        return height, width

    def draw(self):
        canvas = self.canv
        canvas.saveState()
        canvas.rotate(90)
        canvas.translate(0, -self._rotated_width)
        Paragraph.draw(self)
        canvas.restoreState()


class IsAdminUserRole(permissions.BasePermission):
    def has_permission(self, request, view):
        return bool(request.user and request.user.is_authenticated and request.user.role == request.user.Roles.ADMIN)


def _percent(numerator: int, denominator: int) -> float:
    return round((numerator / denominator) * 100, 1) if denominator else 0.0


def _format_percent(value: float, lang: str) -> str:
    text = f"{value:.1f}".rstrip("0").rstrip(".")
    if lang in {"ru", "kg"}:
        text = text.replace(".", ",")
    return f"{text}%"


def _count_with_percent(count: int, denominator: int, lang: str) -> str:
    return f"{count} / {_format_percent(_percent(count, denominator), lang)}" if denominator else f"{count} / {_format_percent(0, lang)}"


def _apply_profile_filters(qs: QuerySet[AlumniProfile], request) -> QuerySet[AlumniProfile]:
    exact_filters = {
        "graduation_year": "graduation_year",
        "academic_group": "academic_group_id",
        "study_form": "study_form",
        "degree_level": "degree_level",
        "employment_status": "employment_status",
        "direction_code": "academic_group__direction_code",
    }
    for query_name, orm_name in exact_filters.items():
        value = request.query_params.get(query_name)
        if value:
            qs = qs.filter(**{orm_name: value})

    search = request.query_params.get("search")
    if search:
        qs = qs.filter(
            Q(user__first_name__icontains=search)
            | Q(user__last_name__icontains=search)
            | Q(user__email__icontains=search)
            | Q(user__username__icontains=search)
            | Q(academic_group__name__icontains=search)
            | Q(direction__icontains=search)
            | Q(profile__icontains=search)
            | Q(specialty__icontains=search)
            | Q(workplace__icontains=search)
            | Q(position__icontains=search)
        )
    return qs


@extend_schema(responses={200: dict})
class EmploymentStatsView(APIView):
    permission_classes = (permissions.AllowAny,)

    def get(self, request):
        base_qs = _apply_profile_filters(AlumniProfile.objects.all(), request)
        qs = base_qs.values("graduation_year").annotate(
            total=Count("id"),
            surveyed=Count("id", filter=Q(is_surveyed=True)),
            employed_specialty=Count("id", filter=Q(employment_status=AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY)),
            employed_not_specialty=Count("id", filter=Q(employment_status=AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY)),
            self_employed=Count("id", filter=Q(employment_status=AlumniProfile.EmploymentStatus.SELF_EMPLOYED)),
            continuing_education=Count(
                "id",
                filter=(
                    Q(employment_status=AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION)
                    | (Q(continuing_education_place__isnull=False) & ~Q(continuing_education_place=""))
                ),
            ),
            unemployed=Count("id", filter=Q(employment_status=AlumniProfile.EmploymentStatus.UNEMPLOYED)),
            lost_contact=Count("id", filter=Q(employment_status=AlumniProfile.EmploymentStatus.LOST_CONTACT)),
        ).order_by("graduation_year")

        result = {}
        status_distribution = defaultdict(int)
        for entry in qs:
            year = entry["graduation_year"] or "Без года"
            total = entry["total"]
            surveyed = entry["surveyed"] or total
            employed = entry["employed_specialty"] + entry["employed_not_specialty"] + entry["self_employed"]
            status_distribution["employed_specialty"] += entry["employed_specialty"]
            status_distribution["employed_not_specialty"] += entry["employed_not_specialty"]
            status_distribution["self_employed"] += entry["self_employed"]
            status_distribution["continuing_education"] += entry["continuing_education"]
            status_distribution["unemployed"] += entry["unemployed"]
            status_distribution["lost_contact"] += entry["lost_contact"]
            status_distribution["unknown"] += entry["lost_contact"]

            result[year] = {
                "total": total,
                "surveyed": surveyed,
                "employed": employed,
                "unemployed": entry["unemployed"],
                "employed_specialty": entry["employed_specialty"],
                "employed_not_specialty": entry["employed_not_specialty"],
                "self_employed": entry["self_employed"],
                "continuing_education": entry["continuing_education"],
                "lost_contact": entry["lost_contact"],
                "unknown": entry["lost_contact"],
                "percent_employed": _percent(employed, surveyed),
                "percent_employed_specialty": _percent(entry["employed_specialty"], surveyed),
                "percent_employed_not_specialty": _percent(entry["employed_not_specialty"], surveyed),
                "percent_self_employed": _percent(entry["self_employed"], surveyed),
                "percent_continuing_education": _percent(entry["continuing_education"], surveyed),
                "percent_unemployed": _percent(entry["unemployed"], surveyed),
                "percent_lost_contact": _percent(entry["lost_contact"], total),
                "percent_unknown": _percent(entry["lost_contact"], total),
            }

        result["meta"] = {
            "total_employers": Employer.objects.count(),
            "status_distribution": dict(status_distribution),
        }
        return Response(result)


@extend_schema(responses={200: OpenApiTypes.BINARY})
class EmploymentReportExportView(APIView):
    permission_classes = (IsAdminUserRole,)
    throttle_scope = "reports"

    def get(self, request, export_format: str = "pdf"):
        export_format = export_format.lower().strip()
        if export_format == "doc":
            export_format = "docx"
        if export_format == "excel":
            export_format = "xlsx"

        report_language = self._report_language(request)
        profiles = self._filtered_profiles(request)
        selection_description = self._selection_description(request, report_language)

        if export_format == "pdf":
            return self._pdf_response(profiles, selection_description, report_language)
        if export_format == "docx":
            return self._docx_response(profiles, selection_description, report_language)
        if export_format == "xlsx":
            return self._xlsx_response(profiles, selection_description, report_language)

        return Response(
            {"detail": _texts(report_language)["unsupported_export"]},
            status=status.HTTP_400_BAD_REQUEST,
        )

    def _filtered_profiles(self, request) -> list[AlumniProfile]:
        qs = AlumniProfile.objects.select_related("user", "employer", "academic_group").order_by(
            "academic_group__graduation_year",
            "academic_group__name",
            "user__last_name",
            "user__first_name",
        )
        return list(_apply_profile_filters(qs, request))

    def _report_language(self, request) -> str:
        explicit_language = request.query_params.get("lang")
        if explicit_language:
            return _normalize_report_language(explicit_language)
        return _normalize_report_language(request.headers.get("Accept-Language"))

    def _report_filename(self, extension: str, lang: str) -> str:
        stamp = timezone.localtime().strftime("%Y-%m-%d_%H-%M")
        return f"{_texts(lang)['file_stem']}_{stamp}.{extension}"

    def _ascii_filename(self, filename: str) -> str:
        fallback = filename.encode("ascii", "ignore").decode("ascii")
        fallback = re.sub(r"[^A-Za-z0-9_.-]+", "_", fallback).strip("._-")
        return fallback or "employment_report"

    def _attachment_header(self, filename: str) -> str:
        fallback = self._ascii_filename(filename)
        return f"attachment; filename=\"{fallback}\"; filename*=UTF-8''{quote(filename)}"

    def _selection_description(self, request, lang: str) -> str:
        labels = _texts(lang)
        parts = []
        year = request.query_params.get("graduation_year")
        group_id = request.query_params.get("academic_group")
        study_form = request.query_params.get("study_form")
        degree_level = request.query_params.get("degree_level")
        status_value = request.query_params.get("employment_status")
        direction_code = request.query_params.get("direction_code")
        search = request.query_params.get("search")
        if year:
            parts.append(f"{labels['year_filter']}: {year}")
        if group_id:
            group = AcademicGroup.objects.filter(pk=group_id).first()
            parts.append(f"{labels['group_filter']}: {group.name if group else group_id}")
        if direction_code:
            parts.append(f"{labels['direction_filter']}: {direction_code}")
        if study_form:
            parts.append(f"{labels['study_form_filter']}: {_choice_label('study_form', study_form, lang)}")
        if degree_level:
            parts.append(f"{labels['degree_level_filter']}: {_choice_label('degree_level', degree_level, lang)}")
        if status_value:
            parts.append(f"{labels['status_filter']}: {_choice_label('employment_status', status_value, lang)}")
        if search:
            parts.append(f"{labels['search_filter']}: {search}")
        return " · ".join(parts) if parts else labels["all_selection"]

    def _register_font(self) -> str:
        paths = [
            os.getenv("PDF_FONT_PATH"),
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        ]
        for path in paths:
            if path and Path(path).exists():
                pdfmetrics.registerFont(TTFont("ReportFont", path))
                return "ReportFont"
        return "Helvetica"

    def _styles(self, font_name: str):
        base = getSampleStyleSheet()
        base.add(ParagraphStyle(name="ReportTitle", fontName=font_name, fontSize=12, leading=15, alignment=TA_CENTER, spaceAfter=2 * mm))
        base.add(ParagraphStyle(name="ReportSubTitle", fontName=font_name, fontSize=8.5, leading=10.5, alignment=TA_CENTER, spaceAfter=1.5 * mm))
        base.add(ParagraphStyle(name="ReportMeta", fontName=font_name, fontSize=8, leading=10, alignment=TA_CENTER, spaceAfter=3 * mm))
        base.add(ParagraphStyle(name="ReportSection", fontName=font_name, fontSize=10.5, leading=12.5, alignment=TA_CENTER, spaceBefore=1 * mm, spaceAfter=3 * mm))
        base.add(ParagraphStyle(name="Cell", fontName=font_name, fontSize=5.6, leading=6.4, alignment=TA_LEFT))
        base.add(ParagraphStyle(name="CellCenter", fontName=font_name, fontSize=5.6, leading=6.4, alignment=TA_CENTER))
        base.add(ParagraphStyle(name="HeaderCell", fontName=font_name, fontSize=5.8, leading=6.5, alignment=TA_CENTER))
        base.add(ParagraphStyle(name="HeaderCellSmall", fontName=font_name, fontSize=5.3, leading=5.9, alignment=TA_CENTER))
        base.add(ParagraphStyle(name="VerticalHeader", fontName=font_name, fontSize=5.4, leading=6.0, alignment=TA_CENTER))
        base["Normal"].fontName = font_name
        base["Normal"].fontSize = 8
        base["Normal"].leading = 10
        return base

    def _p(self, value, styles, center: bool = False, style_name: str | None = None) -> Paragraph:
        text = "—" if value in (None, "") else str(value)
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        selected_style = style_name or ("CellCenter" if center else "Cell")
        return Paragraph(text, styles[selected_style])

    def _vertical_p(self, value, styles) -> RotatedParagraph:
        text = "—" if value in (None, "") else str(value)
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        return RotatedParagraph(text, styles["VerticalHeader"])

    def _work_text(self, profile: AlumniProfile) -> str:
        company = profile.employer.company_name if profile.employer else profile.workplace
        parts = [part for part in (company, profile.position) if part]
        return "\n".join(parts) or "—"

    def _not_working_text(self, profile: AlumniProfile, lang: str) -> str:
        labels = _texts(lang)
        if profile.employment_status == AlumniProfile.EmploymentStatus.UNEMPLOYED:
            return labels["not_working"]
        if profile.employment_status == AlumniProfile.EmploymentStatus.LOST_CONTACT:
            return labels["unknown"]
        return "—"

    def _full_name(self, profile: AlumniProfile) -> str:
        return profile.user.get_full_name() or profile.user.username

    def _full_name_with_group(self, profile: AlumniProfile) -> str:
        full_name = self._full_name(profile)
        group_name = profile.academic_group.name if profile.academic_group else ""
        return f"{full_name} ({group_name})" if group_name else full_name

    def _direction_profile_text(self, profile: AlumniProfile) -> str:
        direction = profile.direction or profile.specialty
        group_direction = profile.academic_group.direction_name if profile.academic_group else ""
        profile_name = profile.profile or (profile.academic_group.profile if profile.academic_group else "")
        return "\n".join(part for part in (direction or group_direction, profile_name) if part) or "—"

    def _graduate_headers(self, lang: str) -> list[str]:
        labels = _texts(lang)
        return [
            labels["number"],
            labels["full_name"],
            labels["group"],
            labels["year_graduation"],
            labels["specialty_diploma"],
            labels["employed_specialty"],
            labels["not_specialty_self_employed"],
            labels["not_working_unknown"],
            labels["continued_education"],
            labels["useful_work"],
            labels["self_study"],
        ]

    def _profile_row_values(self, profile: AlumniProfile, index: int, lang: str) -> list[str | int | None]:
        status_value = profile.employment_status
        work = self._work_text(profile)
        return [
            index,
            self._full_name(profile),
            profile.academic_group.name if profile.academic_group else "—",
            profile.graduation_year or "—",
            profile.specialty or self._direction_profile_text(profile),
            work if status_value == AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY else "—",
            work if status_value in (AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY, AlumniProfile.EmploymentStatus.SELF_EMPLOYED) else "—",
            self._not_working_text(profile, lang),
            profile.continuing_education_place
            if status_value == AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION or profile.continuing_education_place
            else "—",
            profile.useful_subjects or "—",
            profile.self_study_topics or "—",
        ]

    def _graduate_main_headers(self, lang: str) -> list[str]:
        labels = _texts(lang)
        return [
            labels["number"],
            labels["full_name"],
            labels["group"],
            labels["graduation_year"],
            labels["specialty"],
            labels["employed_specialty"],
            labels["not_specialty_self_employed"],
            labels["not_working_unknown"],
            labels["continued_education"],
        ]

    def _profile_main_row_values(self, profile: AlumniProfile, index: int, lang: str) -> list[str | int | None]:
        status_value = profile.employment_status
        work = self._work_text(profile)
        return [
            index,
            self._full_name(profile),
            profile.academic_group.name if profile.academic_group else "—",
            profile.graduation_year or "—",
            profile.specialty or self._direction_profile_text(profile),
            work if status_value == AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY else "—",
            work if status_value in (AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY, AlumniProfile.EmploymentStatus.SELF_EMPLOYED) else "—",
            self._not_working_text(profile, lang),
            profile.continuing_education_place
            if status_value == AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION or profile.continuing_education_place
            else "—",
        ]

    def _survey_headers(self, lang: str) -> list[str]:
        labels = _texts(lang)
        return [labels["number"], labels["full_name"], labels["useful_work"], labels["self_study"]]

    def _profile_survey_row_values(self, profile: AlumniProfile, index: int) -> list[str | int | None]:
        return [
            index,
            self._full_name(profile),
            profile.useful_subjects or "—",
            profile.self_study_topics or "—",
        ]

    def _graduates_main_table(self, profiles: Iterable[AlumniProfile], styles, lang: str) -> Table:
        data = [[self._p(value, styles, center=True) for value in self._graduate_main_headers(lang)]]
        for index, profile in enumerate(profiles, start=1):
            row = self._profile_main_row_values(profile, index, lang)
            data.append([self._p(value, styles, center=column_index in (0, 2, 3, 4)) for column_index, value in enumerate(row)])

        table = Table(
            data,
            repeatRows=1,
            colWidths=[8 * mm, 36 * mm, 21 * mm, 13 * mm, 23 * mm, 47 * mm, 44 * mm, 31 * mm, 50 * mm],
        )
        table.setStyle(self._table_style())
        return table

    def _survey_table(self, profiles: Iterable[AlumniProfile], styles, lang: str) -> Table:
        data = [[self._p(value, styles, center=True) for value in self._survey_headers(lang)]]
        for index, profile in enumerate(profiles, start=1):
            row = self._profile_survey_row_values(profile, index)
            data.append([self._p(value, styles, center=column_index == 0) for column_index, value in enumerate(row)])

        table = Table(
            data,
            repeatRows=1,
            colWidths=[8 * mm, 48 * mm, 112 * mm, 105 * mm],
        )
        table.setStyle(self._table_style())
        return table

    def _graduate_reference_headers(self, styles, lang: str) -> list[list[Paragraph | RotatedParagraph | str]]:
        labels = _texts(lang)
        return [
            [
                self._p(labels["number"], styles, center=True, style_name="HeaderCell"),
                self._p(labels["full_name"], styles, center=True, style_name="HeaderCell"),
                self._vertical_p(labels["year_graduation"], styles),
                self._vertical_p(labels["specialty_diploma"], styles),
                self._p(labels["employment"], styles, center=True, style_name="HeaderCell"),
                "",
                "",
                self._p(labels["continued_education_where"], styles, center=True, style_name="HeaderCell"),
                self._p(labels["useful_question"], styles, center=True, style_name="HeaderCellSmall"),
                self._p(labels["self_study_question"], styles, center=True, style_name="HeaderCellSmall"),
            ],
            [
                "",
                "",
                "",
                "",
                self._p(labels["employed_specialty_header"], styles, center=True, style_name="HeaderCellSmall"),
                self._p(labels["not_specialty_self_employed_header"], styles, center=True, style_name="HeaderCellSmall"),
                self._p(labels["not_working_unknown_header"], styles, center=True, style_name="HeaderCellSmall"),
                "",
                "",
                "",
            ],
        ]

    def _pdf_profile_row_values(self, profile: AlumniProfile, index: int, lang: str) -> list[str | int | None]:
        status_value = profile.employment_status
        work = self._work_text(profile)
        return [
            index,
            self._full_name_with_group(profile),
            profile.graduation_year or "—",
            profile.specialty or self._direction_profile_text(profile),
            work if status_value == AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY else "—",
            work if status_value in (AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY, AlumniProfile.EmploymentStatus.SELF_EMPLOYED) else "—",
            self._not_working_text(profile, lang),
            profile.continuing_education_place
            if status_value == AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION or profile.continuing_education_place
            else "—",
            profile.useful_subjects or "—",
            profile.self_study_topics or "—",
        ]

    def _graduates_table(self, profiles: Iterable[AlumniProfile], styles, lang: str) -> Table:
        data = self._graduate_reference_headers(styles, lang)
        for index, profile in enumerate(profiles, start=1):
            row = self._pdf_profile_row_values(profile, index, lang)
            data.append([self._p(value, styles, center=column_index in (0, 2, 3)) for column_index, value in enumerate(row)])

        table = Table(
            data,
            repeatRows=2,
            colWidths=[7 * mm, 46 * mm, 10 * mm, 13 * mm, 35 * mm, 31 * mm, 18 * mm, 34 * mm, 55 * mm, 40 * mm],
            rowHeights=[31 * mm, 13 * mm] + [None] * max(len(data) - 2, 0),
        )
        table.setStyle(self._graduate_table_style())
        return table

    def _summary_headers(self, lang: str) -> list[str]:
        labels = _texts(lang)
        return [
            labels["graduation_year"],
            labels["direction_profile"],
            labels["form_level"],
            labels["graduates_count"],
            labels["surveyed_count"],
            labels["works_specialty_percent"],
            labels["works_not_specialty_percent"],
            labels["self_employed_percent"],
            labels["continued_education_percent"],
            labels["not_working_percent"],
            labels["unknown_percent"],
        ]

    def _summary_row_values(self, profiles: Iterable[AlumniProfile], lang: str) -> list[list[str | int]]:
        buckets: dict[tuple[str, str, str, str, int | None], list[AlumniProfile]] = defaultdict(list)
        for profile in profiles:
            group = profile.academic_group
            direction_code = group.direction_code if group else "—"
            profile_name = profile.profile or (group.profile if group else "—")
            study_form_value = profile.study_form or (group.study_form if group else None)
            degree_level_value = profile.degree_level or (group.degree_level if group else None)
            study_form = _choice_label("study_form", study_form_value, lang)
            degree_level = _choice_label("degree_level", degree_level_value, lang)
            buckets[(direction_code, profile_name, study_form, degree_level, profile.graduation_year)].append(profile)

        rows: list[list[str | int]] = []
        for (direction_code, profile_name, study_form, degree_level, year), items in sorted(
            buckets.items(),
            key=lambda item: (item[0][4] or 0, item[0][0], item[0][1], item[0][2], item[0][3]),
        ):
            total = len(items)
            surveyed = sum(1 for item in items if item.is_surveyed)
            denominator = surveyed or total
            counts = defaultdict(int)
            continuing_count = 0
            for item in items:
                counts[item.employment_status] += 1
                if item.employment_status == AlumniProfile.EmploymentStatus.CONTINUING_EDUCATION or item.continuing_education_place:
                    continuing_count += 1
            direction_label = f"{direction_code}\n{profile_name}"
            rows.append([
                year or _texts(lang)["no_year"],
                direction_label,
                f"{study_form}\n{degree_level}",
                total,
                surveyed,
                _count_with_percent(counts[AlumniProfile.EmploymentStatus.EMPLOYED_SPECIALTY], denominator, lang),
                _count_with_percent(counts[AlumniProfile.EmploymentStatus.EMPLOYED_NOT_SPECIALTY], denominator, lang),
                _count_with_percent(counts[AlumniProfile.EmploymentStatus.SELF_EMPLOYED], denominator, lang),
                _count_with_percent(continuing_count, denominator, lang),
                _count_with_percent(counts[AlumniProfile.EmploymentStatus.UNEMPLOYED], denominator, lang),
                _count_with_percent(counts[AlumniProfile.EmploymentStatus.LOST_CONTACT], total, lang),
            ])
        return rows

    def _summary_table(self, profiles: Iterable[AlumniProfile], styles, lang: str) -> Table:
        data = [[self._p(value, styles, center=True) for value in self._summary_headers(lang)]]
        for row in self._summary_row_values(profiles, lang):
            data.append([self._p(value, styles, center=True) for value in row])

        table = Table(
            data,
            repeatRows=1,
            colWidths=[18 * mm, 45 * mm, 28 * mm, 18 * mm, 18 * mm, 26 * mm, 27 * mm, 22 * mm, 25 * mm, 22 * mm, 22 * mm],
        )
        table.setStyle(self._table_style())
        return table

    def _text_value(self, value) -> str:
        return "—" if value in (None, "") else str(value)

    def _pdf_response(self, profiles: list[AlumniProfile], selection_description: str, lang: str) -> HttpResponse:
        buffer = BytesIO()
        font_name = self._register_font()
        styles = self._styles(font_name)
        labels = _texts(lang)

        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            leftMargin=4 * mm,
            rightMargin=4 * mm,
            topMargin=6 * mm,
            bottomMargin=6 * mm,
            title=labels["title"],
        )
        story = [
            Paragraph(labels["title"], styles["ReportTitle"]),
            Paragraph(labels["subtitle"], styles["ReportSubTitle"]),
            Paragraph(selection_description, styles["ReportMeta"]),
            Paragraph(labels["section_1"], styles["ReportSection"]),
            self._graduates_table(profiles, styles, lang),
            PageBreak(),
            Paragraph(labels["title"], styles["ReportTitle"]),
            Paragraph(labels["subtitle"], styles["ReportSubTitle"]),
            Paragraph(selection_description, styles["ReportMeta"]),
            Paragraph(labels["section_2"], styles["ReportSection"]),
            self._summary_table(profiles, styles, lang),
            Spacer(1, 12 * mm),
            Paragraph(labels["chair_signature"], styles["Normal"]),
        ]

        doc.build(story)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = self._attachment_header(self._report_filename("pdf", lang))
        return response

    def _docx_response(self, profiles: list[AlumniProfile], selection_description: str, lang: str) -> HttpResponse:
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt
        except ImportError as exc:
            return Response(
                {"detail": f"DOCX export dependency is not installed: {exc}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        def set_landscape(section):
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            section.left_margin = Cm(0.7)
            section.right_margin = Cm(0.7)
            section.top_margin = Cm(0.7)
            section.bottom_margin = Cm(0.7)

        def set_cell_margins(cell, margin_twips: str = "45") -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_name in ("top", "left", "bottom", "right"):
                margin = tc_mar.find(qn(f"w:{margin_name}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_name}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), margin_twips)
                margin.set(qn("w:type"), "dxa")

        def set_cell_width(cell, width_cm: float) -> None:
            cell.width = Cm(width_cm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width_cm * 567)))
            tc_w.set(qn("w:type"), "dxa")

        def set_cell_text_direction(cell, direction: str = "btLr") -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            text_direction = tc_pr.find(qn("w:textDirection"))
            if text_direction is None:
                text_direction = OxmlElement("w:textDirection")
                tc_pr.append(text_direction)
            text_direction.set(qn("w:val"), direction)

        def set_repeat_header(row) -> None:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                tbl_header = OxmlElement("w:tblHeader")
                tbl_header.set(qn("w:val"), "true")
                tr_pr.append(tbl_header)

        def set_fixed_table_layout(table, width_cm: float | None = None) -> None:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            tbl_pr = table._tbl.tblPr
            tbl_layout = tbl_pr.find(qn("w:tblLayout"))
            if tbl_layout is None:
                tbl_layout = OxmlElement("w:tblLayout")
                tbl_pr.append(tbl_layout)
            tbl_layout.set(qn("w:type"), "fixed")
            if width_cm is not None:
                tbl_w = tbl_pr.find(qn("w:tblW"))
                if tbl_w is None:
                    tbl_w = OxmlElement("w:tblW")
                    tbl_pr.append(tbl_w)
                tbl_w.set(qn("w:type"), "dxa")
                tbl_w.set(qn("w:w"), str(int(width_cm * 567)))

        def apply_column_widths(table, widths: list[float]) -> None:
            for column, width in zip(table.columns, widths):
                column.width = Cm(width)
            for row in table.rows:
                for index, width in enumerate(widths):
                    set_cell_width(row.cells[index], width)

        def set_row_height(row, height_cm: float, exact: bool = False) -> None:
            row.height = Cm(height_cm)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST

        def set_cell_text(
            cell,
            value,
            *,
            bold: bool = False,
            align_center: bool = False,
            font_size: float = 7.2,
            vertical_middle: bool = False,
            rotate: bool = False,
        ) -> None:
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for line_index, line in enumerate(self._text_value(value).splitlines() or ["—"]):
                if line_index:
                    paragraph.add_run().add_break()
                run = paragraph.add_run(line)
                run.bold = bold
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if vertical_middle else WD_CELL_VERTICAL_ALIGNMENT.TOP
            if rotate:
                set_cell_text_direction(cell)
            set_cell_margins(cell)

        def add_centered_paragraph(text: str, *, size: float, bold: bool = False, spacing_after: float = 2) -> None:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(spacing_after)
            run = paragraph.add_run(text)
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        labels = _texts(lang)

        document = Document()
        set_landscape(document.sections[0])
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(8)
        normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        add_centered_paragraph(labels["title"], size=12, bold=True, spacing_after=1)
        add_centered_paragraph(labels["subtitle"], size=9, spacing_after=1)
        add_centered_paragraph(selection_description, size=8, spacing_after=4)
        add_centered_paragraph(labels["section_1"], size=10.5, bold=True, spacing_after=4)

        graduate_widths = [1.15, 4.2, 0.9, 1.2, 3.25, 2.85, 1.65, 3.2, 5.15, 3.6]
        graduate_table = document.add_table(rows=2, cols=10)
        graduate_table.style = "Table Grid"
        set_fixed_table_layout(graduate_table, sum(graduate_widths))
        apply_column_widths(graduate_table, graduate_widths)
        set_row_height(graduate_table.rows[0], 2.25, exact=True)
        set_row_height(graduate_table.rows[1], 1.05, exact=True)
        set_repeat_header(graduate_table.rows[0])
        set_repeat_header(graduate_table.rows[1])

        header_0 = graduate_table.rows[0].cells
        header_1 = graduate_table.rows[1].cells
        header_0[0].merge(header_1[0])
        header_0[1].merge(header_1[1])
        header_0[2].merge(header_1[2])
        header_0[3].merge(header_1[3])
        header_0[4].merge(header_0[6])
        header_0[7].merge(header_1[7])
        header_0[8].merge(header_1[8])
        header_0[9].merge(header_1[9])

        set_cell_text(graduate_table.cell(0, 0), labels["number"], bold=True, align_center=True, font_size=7.6, vertical_middle=True)
        set_cell_text(graduate_table.cell(0, 1), labels["full_name"], bold=True, align_center=True, font_size=7.6, vertical_middle=True)
        set_cell_text(graduate_table.cell(0, 2), labels["year_graduation"], bold=True, align_center=True, font_size=7.2, vertical_middle=True, rotate=True)
        set_cell_text(graduate_table.cell(0, 3), labels["specialty_diploma"], bold=True, align_center=True, font_size=6.8, vertical_middle=True, rotate=True)
        set_cell_text(graduate_table.cell(0, 4), labels["employment"], bold=True, align_center=True, font_size=7.6, vertical_middle=True)
        set_cell_text(graduate_table.cell(1, 4), labels["employed_specialty_header"], bold=True, align_center=True, font_size=6.8, vertical_middle=True)
        set_cell_text(graduate_table.cell(1, 5), labels["not_specialty_self_employed_header"], bold=True, align_center=True, font_size=6.8, vertical_middle=True)
        set_cell_text(graduate_table.cell(1, 6), labels["not_working_unknown_header"], bold=True, align_center=True, font_size=6.8, vertical_middle=True)
        set_cell_text(graduate_table.cell(0, 7), labels["continued_education_where"], bold=True, align_center=True, font_size=7.0, vertical_middle=True)
        set_cell_text(graduate_table.cell(0, 8), labels["useful_question"], bold=True, align_center=True, font_size=6.6, vertical_middle=True)
        set_cell_text(graduate_table.cell(0, 9), labels["self_study_question"], bold=True, align_center=True, font_size=6.6, vertical_middle=True)

        for index, profile in enumerate(profiles, start=1):
            cells = graduate_table.add_row().cells
            for column_index, value in enumerate(self._pdf_profile_row_values(profile, index, lang)):
                set_cell_width(cells[column_index], graduate_widths[column_index])
                set_cell_text(
                    cells[column_index],
                    value,
                    align_center=column_index in (0, 2, 3),
                    font_size=6.7,
                    vertical_middle=column_index in (0, 2, 3),
                )

        document.add_page_break()
        add_centered_paragraph(labels["title"], size=12, bold=True, spacing_after=1)
        add_centered_paragraph(labels["subtitle"], size=9, spacing_after=1)
        add_centered_paragraph(selection_description, size=8, spacing_after=4)
        add_centered_paragraph(labels["section_2"], size=10.5, bold=True, spacing_after=4)

        summary_headers = self._summary_headers(lang)
        summary_rows = self._summary_row_values(profiles, lang)
        summary_widths = [1.7, 4.2, 2.6, 1.7, 1.7, 2.45, 2.55, 2.05, 2.35, 2.05, 2.05]
        summary_table = document.add_table(rows=1, cols=len(summary_headers))
        summary_table.style = "Table Grid"
        set_fixed_table_layout(summary_table, sum(summary_widths))
        apply_column_widths(summary_table, summary_widths)
        set_repeat_header(summary_table.rows[0])
        for cell, header, width in zip(summary_table.rows[0].cells, summary_headers, summary_widths):
            set_cell_width(cell, width)
            set_cell_text(cell, header, bold=True, align_center=True, font_size=6.7, vertical_middle=True)
        for row in summary_rows:
            cells = summary_table.add_row().cells
            for column_index, value in enumerate(row):
                set_cell_width(cells[column_index], summary_widths[column_index])
                set_cell_text(cells[column_index], value, align_center=True, font_size=6.7, vertical_middle=True)

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(18)
        run = paragraph.add_run(labels["chair_signature"])
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = self._attachment_header(self._report_filename("docx", lang))
        return response

    def _xlsx_response(self, profiles: list[AlumniProfile], selection_description: str, lang: str) -> HttpResponse:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
            from openpyxl.utils import get_column_letter
        except ImportError as exc:
            return Response(
                {"detail": f"XLSX export dependency is not installed: {exc}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        labels = _texts(lang)
        workbook = Workbook()
        border_side = Side(style="thin", color="111827")
        border = Border(left=border_side, right=border_side, top=border_side, bottom=border_side)
        header_font = Font(name="Arial", size=10, bold=True, color="111827")
        title_font = Font(name="Arial", size=12, bold=True, color="111827")
        body_font = Font(name="Arial", size=10, color="111827")
        header_fill = PatternFill(fill_type="solid", fgColor="F8FAFC")

        def prepare_sheet(sheet, title: str, headers: list[str], rows: list[list[str | int | None]], widths: list[int]):
            sheet.title = title
            sheet.append([labels["title"]])
            sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
            sheet.cell(1, 1).font = title_font
            sheet.cell(1, 1).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            sheet.append([labels["subtitle"]])
            sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
            sheet.cell(2, 1).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            sheet.append([selection_description])
            sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=len(headers))
            sheet.cell(3, 1).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            sheet.append(headers)
            for cell in sheet[4]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = border
            for row in rows:
                sheet.append([self._text_value(value) for value in row])
            for row in sheet.iter_rows(min_row=5, max_row=sheet.max_row, min_col=1, max_col=len(headers)):
                for cell in row:
                    cell.font = body_font
                    cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                    cell.border = border
            sheet.freeze_panes = "A5"
            sheet.auto_filter.ref = f"A4:{get_column_letter(len(headers))}{sheet.max_row}"
            for index, width in enumerate(widths, start=1):
                sheet.column_dimensions[get_column_letter(index)].width = width
            sheet.row_dimensions[1].height = 34
            sheet.row_dimensions[2].height = 28
            sheet.row_dimensions[3].height = 24

        graduate_headers = self._graduate_headers(lang)
        graduate_rows = [self._profile_row_values(profile, index, lang) for index, profile in enumerate(profiles, start=1)]
        prepare_sheet(
            workbook.active,
            labels["sheet_graduates"],
            graduate_headers,
            graduate_rows,
            [7, 30, 18, 14, 24, 36, 34, 24, 30, 40, 34],
        )

        summary_headers = self._summary_headers(lang)
        summary_rows = self._summary_row_values(profiles, lang)
        summary_sheet = workbook.create_sheet(labels["sheet_summary"])
        prepare_sheet(
            summary_sheet,
            labels["sheet_summary"],
            summary_headers,
            summary_rows,
            [14, 40, 24, 20, 20, 22, 24, 20, 24, 20, 20],
        )

        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = self._attachment_header(self._report_filename("xlsx", lang))
        return response

    def _graduate_table_style(self) -> TableStyle:
        return TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#111827")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("VALIGN", (0, 2), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 1), "CENTER"),
            ("SPAN", (0, 0), (0, 1)),
            ("SPAN", (1, 0), (1, 1)),
            ("SPAN", (2, 0), (2, 1)),
            ("SPAN", (3, 0), (3, 1)),
            ("SPAN", (4, 0), (6, 0)),
            ("SPAN", (7, 0), (7, 1)),
            ("SPAN", (8, 0), (8, 1)),
            ("SPAN", (9, 0), (9, 1)),
            ("LEFTPADDING", (0, 0), (-1, -1), 1.2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 1.2),
            ("TOPPADDING", (0, 0), (-1, -1), 1.4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1.4),
        ])

    def _table_style(self) -> TableStyle:
        return TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#374151")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ])